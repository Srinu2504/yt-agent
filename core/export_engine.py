import os
import io

BASE_DIR    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EXPORTS_DIR = os.path.join(BASE_DIR, "data", "exports")
os.makedirs(EXPORTS_DIR, exist_ok=True)


def export_txt(content: str, filename: str) -> str:
    path = os.path.join(EXPORTS_DIR, f"{filename}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def export_pdf(content: str, title: str, filename: str) -> str:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        raise ImportError("reportlab not installed: pip install reportlab")

    path = os.path.join(EXPORTS_DIR, f"{filename}.pdf")
    buf  = io.BytesIO()
    doc  = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm,  bottomMargin=2.5*cm,
    )

    styles     = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "t", parent=styles["Title"],
        fontSize=18, spaceAfter=16,
        textColor=colors.HexColor("#1a1a2e")
    )
    body_style = ParagraphStyle(
        "b", parent=styles["Normal"],
        fontSize=11, leading=17, spaceAfter=8
    )
    heading_style = ParagraphStyle(
        "h", parent=styles["Heading2"],
        fontSize=13, spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor("#0066cc")
    )

    story = [Paragraph(title, title_style), Spacer(1, 0.3*cm)]

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
        elif line.startswith("## "):
            safe = line[3:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe, heading_style))
        elif line.startswith("# "):
            safe = line[2:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe, title_style))
        elif line.startswith("- ") or line.startswith("* "):
            safe = line[2:].replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(f"• {safe}", body_style))
        else:
            safe = line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(safe, body_style))

    doc.build(story)

    with open(path, "wb") as f:
        f.write(buf.getvalue())

    return path


def export_docx(content: str, title: str, filename: str) -> str:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed: pip install python-docx")

    path = os.path.join(EXPORTS_DIR, f"{filename}.docx")
    doc  = Document()

    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    doc.save(path)
    return path


def get_export_bytes(path: str) -> bytes:
    if os.path.exists(path):
        with open(path, "rb") as f:
            return f.read()
    return b""
